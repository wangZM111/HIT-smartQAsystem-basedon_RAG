import pdfplumber
import embedding_json
import json
import doubao_API as API
import time
from docx import Document
import os
def read_pdf(pdf_filepath):
    all_text = []
    with pdfplumber.open(pdf_filepath) as pdf:
        for page in pdf.pages:
            text = page.extract_text()
            if text:  # 有的页可能没有可提取文本
                all_text.append(text)
    # 合并所有页文本
    text = '\n'.join(all_text)
    api_response = API.topic_generation(text)
    current_time = time.localtime()  # 获取本地时间
    formatted_date = time.strftime("%Y/%m/%d", current_time)
    url = f"{api_response}/{formatted_date}"
    title  = api_response
    content = text
    data_list = [{
        "url": url,
        "title": title,
        "content": content
    }]
    safe_url = f"{api_response}_{formatted_date}".replace('/', '_').replace('\\', '_')
    current_dir = os.path.dirname(os.path.abspath(__file__))
    json_path = os.path.join(current_dir, "用户上传文件文件集",f"{safe_url}.json")
    with open(json_path, 'w', encoding='utf-8') as f:
        json.dump(data_list, f, ensure_ascii=False, indent=2)
    embedding_json.process_file_and_save_information(json_path,flag=False)
    return 1

def read_docx(doc_filepath):
    doc = Document(doc_filepath)
    full_text = []
    # 提取段落文字
    for para in doc.paragraphs:
        full_text.append(para.text)
    # 提取表格内容
    for table in doc.tables:
        for row in table.rows:
            row_data = [cell.text.strip() for cell in row.cells]
            full_text.append('\t'.join(row_data))  # 使用制表符连接每行
    text =  '\n'.join(full_text)
    api_response = API.topic_generation(text)
    current_time = time.localtime()  # 获取本地时间
    formatted_date = time.strftime("%Y/%m/%d", current_time)
    url = f"{api_response}/{formatted_date}"
    title = api_response
    content = text
    data_list = [{
        "url": url,
        "title": title,
        "content": content
    }]
    safe_url = f"{api_response}_{formatted_date}".replace('/', '_').replace('\\', '_')
    current_dir = os.path.dirname(os.path.abspath(__file__))
    json_path = os.path.join(current_dir, "用户上传文件文件集",f"{safe_url}.json")
    with open(json_path, 'w', encoding='utf-8') as f:
        json.dump(data_list, f, ensure_ascii=False, indent=2)
    embedding_json.process_file_and_save_information(json_path,flag=False)
    return 1
